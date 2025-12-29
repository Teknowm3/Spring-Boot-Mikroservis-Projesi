#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Converter
Converts proje.md to Proje.docx with all formatting and images
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image

def add_heading(doc, text, level):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text.strip(), level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
    elif level == 2:
        heading.runs[0].font.size = Pt(18)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    elif level == 3:
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 76, 153)
    return heading

def add_paragraph(doc, text, style='Normal'):
    """Add a paragraph with proper formatting"""
    if not text.strip():
        return doc.add_paragraph()
    
    p = doc.add_paragraph(style=style)
    
    # Handle inline formatting (bold, italic, inline code)
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|_.*?_|\[.*?\]\(.*?\))', text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            # Bold
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # Inline code
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78)
        elif part.startswith('_') and part.endswith('_'):
            # Italic
            run = p.add_run(part[1:-1])
            run.italic = True
        elif re.match(r'\[.*?\]\(.*?\)', part):
            # Link (just show as text for now)
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            run = p.add_run(match.group(1))
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
        else:
            # Normal text
            p.add_run(part)
    
    return p

def add_code_block(doc, code, language=''):
    """Add a code block with proper formatting"""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add light gray background effect through shading
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    p._p.get_or_add_pPr().append(shading_elm)
    
    return p

def add_table_from_markdown(doc, lines):
    """Convert markdown table to Word table"""
    # Parse table
    rows = []
    for line in lines:
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            # Skip separator lines
            if not all(re.match(r'^:?-+:?$', cell) for cell in cells):
                rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data
            
            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # Add background color
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '0066CC')
                cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Center align all cells
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Add spacing after table

def add_image(doc, image_path):
    """Add image to document"""
    if not os.path.exists(image_path):
        p = doc.add_paragraph(f"[Image not found: {image_path}]")
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        return
    
    try:
        # Check image size and resize if needed
        with Image.open(image_path) as img:
            width, height = img.size
            max_width = 6.0  # inches
            
            if width > max_width * 96:  # 96 DPI
                aspect_ratio = height / width
                doc.add_picture(image_path, width=Inches(max_width))
            else:
                doc.add_picture(image_path)
        
        # Center the image
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add spacing
    except Exception as e:
        p = doc.add_paragraph(f"[Error loading image {image_path}: {e}]")
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)

def convert_markdown_to_docx(md_file, docx_file):
    """Main conversion function"""
    print(f"Converting {md_file} to {docx_file}...")
    
    # Create document
    doc = Document()
    
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_block = []
    code_language = ''
    in_table = False
    table_lines = []
    
    base_dir = os.path.dirname(md_file)
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line[3:].strip()
                code_block = []
            else:
                in_code_block = False
                add_code_block(doc, '\n'.join(code_block), code_language)
                code_block = []
            i += 1
            continue
        
        if in_code_block:
            code_block.append(line)
            i += 1
            continue
        
        # Tables
        if '|' in line and not in_table:
            in_table = True
            table_lines = [line]
            i += 1
            continue
        
        if in_table:
            if '|' in line:
                table_lines.append(line)
                i += 1
                continue
            else:
                in_table = False
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                continue
        
        # Headings
        if line.startswith('#'):
            level = len(re.match(r'^#+', line).group())
            text = line[level:].strip()
            add_heading(doc, text, min(level, 3))
            i += 1
            continue
        
        # Images
        if line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)
                # Convert relative path to absolute
                if not os.path.isabs(image_path):
                    image_path = os.path.join(base_dir, image_path.replace('/', os.sep))
                
                # Add caption
                if alt_text:
                    p = doc.add_paragraph(alt_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.italic = True
                    p.runs[0].font.size = Pt(10)
                
                add_image(doc, image_path)
            i += 1
            continue
        
        # Horizontal rules
        if re.match(r'^-{3,}$', line):
            doc.add_paragraph('_' * 50)
            i += 1
            continue
        
        # Lists
        if re.match(r'^[\*\-\+]\s+', line) or re.match(r'^\d+\.\s+', line):
            # Bullet or numbered list
            list_item = re.sub(r'^[\*\-\+\d\.]\s+', '', line)
            p = doc.add_paragraph(list_item, style='List Bullet' if re.match(r'^[\*\-\+]', line) else 'List Number')
            i += 1
            continue
        
        # Regular paragraphs
        if line.strip():
            add_paragraph(doc, line)
        else:
            doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Successfully converted to {docx_file}")

if __name__ == '__main__':
    md_file = 'proje.md'
    docx_file = 'Proje.docx'
    
    if not os.path.exists(md_file):
        print(f"❌ Error: {md_file} not found!")
        exit(1)
    
    convert_markdown_to_docx(md_file, docx_file)
    print(f"📄 Word document created: {docx_file}")
